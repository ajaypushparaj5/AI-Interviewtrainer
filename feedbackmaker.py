       
import os
from docx import Document
from docx.shared import Inches

def generate_feedback_doc(report,grade,abnormal_thresholds, strictness=1):
    print("[INFO] Starting feedback document generation...")
    print("[Info] This is the recieved grade:",grade)

    try:
        # Normalize based on strictness level
        if strictness == 2:
            duration_unit = max(1, report["duration_sec"] / 30)
        elif strictness == 3:
            duration_unit = max(1, report["duration_sec"] / 15)
        elif strictness == 4:
            duration_unit = max(1, report["duration_sec"] / 60)
        else:
            duration_unit = 1
        print(f"[INFO] Duration unit calculated: {duration_unit:.2f} seconds")

        norm = lambda count: count / duration_unit

        # Normalize behavior counts
        print("[INFO] Normalizing behavior metrics...")
        normalized_report = {
            "eye_contact_breaks": norm(report["eye_contact_breaks"]),
            "total_blinks": norm(report["total_blinks"]),
            "mouth_touch_count": norm(report["mouth_touch_count"]),
            "nose_touch_count": norm(report["nose_touch_count"]),
            "eye_touch_count": norm(report["eye_touch_count"]),
            "ear_touch_count": norm(report["ear_touch_count"]),
            "neck_touch_count": norm(report["neck_touch_count"]),
            "arms_crossed_for_3_sec_count": norm(report["arms_crossed_for_3_sec_count"]),
            "slouching": norm(report["slouching"]),
            "leg_crossed_count": norm(report["leg_crossed_count"]),
            "leg_bouncing_count": norm(report["leg_bouncing_count"]),
            "hand_on_hip_count": norm(report.get("hand_on_hip_count", 0)),
            "hands_outside_gesture_box_count": norm(report.get("hands_outside_gesture_box_count", 0)),
            "hands_clenched_count": norm(report.get("hands_clenched_count", 0)),
            "hands_behind_back_count": norm(report.get("hands_behind_back_count", 0)),
            "hands_in_pockets_count": norm(report.get("hands_in_pockets_count", 0)),
            "final_bpm": report.get("final_bpm", 0)
        }

        print("[INFO] Normalization complete.")

        
        print("[INFO] Preparing feedback descriptions...")
        
        gesture_feedback = {
        "eye_contact_breaks": {
            "title":"Eye Contact",
            "text": "Avoids eye contact, Glancing at ceiling or floor reflects discomfort or disengagement. Audience feels disconnected; speaker may seem uncertain or underprepared. Practice maintaining soft, steady eye contact with various parts of the audience. ",
            "image": "eye_contact_breaks.png"
            },
        "total_blinks": {
            "title":"Rapid Blinking",
            "text": "Rapid or repeated blinking (more than 20–30 blinks per minute) is often a sign of nervousness, stress, or cognitive overload. It may also indicate discomfort with eye contact or dishonesty, depending on the context. The audience or interviewer may interpret the behavior as a lack of confidence, anxiety, or even evasiveness. It disrupts engagement and reduces perceived trustworthiness. Practice calming techniques such as deep breathing and speaking slower. Improve eye contact comfort by rehearsing with a friend, mirror, or camera. Reduce stress to regulate blinking naturally.",
            "image": "total_blinks.png"
            },
        
        "mouth_touch_count": {
            "title":"Mouth Touching",
            "text" : "Touching the face, mouth, nose, eyes, ears, and neck are soothing gestures that indicate internal discomfort, anxiety, or deception. These displays are common under stress. It can be distracting and potentially lower credibility. Reduces confidence in what’s being said. Increase awareness of this habit. Keep your hands visible and use purposeful gestures instead. Neck / Face Rubbing or Scratching is a Sign of doubt, insecurity, or emotional discomfort. It may weaken trust and show hesitation. Replace with neutral hand positions and affirming verbal cues. ",
            "image": "mouth_touch_count.png",
            "additional": "Fingers in Mouth / Nail Biting are strong pacifying behaviours, often a response to anxiety or lack of self-control. Appears unprofessional and immature. Therefore, it should be avoided. Men tend to massage or stroke their necks to pacify distress. This area is rich with nerves, including the vagus nerve, which, when massaged, will slow down the heart rate. Even a brief touch of the neck will serve to assuage anxiety or discomfort. Neck touching or massaging is a powerful and universal stress reliever and pacifier. It should be avoided."
            },
        
        "nose_touch_count": {
            "title":"Nose Touching",
            "text" : "Touching Face Mouth, Nose, Eyes, Ears and Neck are Soothing gesture that indicates internal discomfort. anxiety, or deception. Common under stress. Can be distracting and signals discomfort, potentially lowering credibility. Reduces confidence in what’s being said. Increase awareness of this habit. Keep hands visible and use purposeful gestures instead.",
            "image": "nose_touch_count.png"
            },
        
        "eye_touch_count": {
            "title":"Eye Touching",
            "text" : "Touching Eyes is a common pacifying behavior, often a response to anxiety or discomfort. It can be distracting and may signal nervousness or dishonesty. Avoid touching your eyes during presentations or interviews to maintain professionalism and credibility. Also, Eyebrow Movement Adds emotional depth; raised brows shows interest and furrowed coneys concern or confusion. Can emphasize or distort your message if mismatched. Use eyebrows expressively but in sync with tone and content.",
            "image": "eye_touch_count.png"
            },
        
        "ear_touch_count": {
            "title":"Ear Touching",
            "text" : "Touching your ears is a common self-soothing gesture, typically triggered by anxiety or discomfort. However, it can be distracting and may unintentionally signal nervousness or a lack of honesty. To maintain a professional and confident impression during presentations or interviews, it's best to avoid this behavior.",
            "image": "ear_touch_count.png"
            },
        
        "neck_touch_count": {
            "title":"Neck Touching",
            "text" : "Scratching the neck is typically interpreted as a pacifying behavior, often linked to self-doubt, stress, or internal conflict. It's a subconscious gesture that may indicate the person is feeling uncertain, uncomfortable, or lacking confidence in what they're saying. In high-stakes settings like interviews or public speaking, this gesture can subtly undermine your message, suggesting hesitation or insincerity. For a more confident and composed presence, it's advisable to become aware of and minimize such self-touch behaviors.",
            "image": "neck_touch_count.png"
            },
        
        "arms_crossed_for_3_sec_count": {
            "title":"Arms Crossed",
            "text": "Both arms are folded together across the chest is a Defensive or protective gesture. Reduces approachability and may signal emotional distancing, a negative attitude, shyness or nervousness. Use open posture with relaxed arms. Crossed Arms with Hands Gripping Arms also conveys Strong emotional discomfort or internal resistance." ,
            "image": "arms_crossed.png"
            },
        
        "slouching": {
            "title":"Slouching",
            "text": "Slouching Indicates lack of composure, energy, or seriousness. Projects low confidence and poor presence. Practice upright posture with an open chest, feet grounded, and shoulders relaxed.",
            "image": "slouching.png",
            "additional": " • A slight forward lean signals engagement. However, leaning back indicates disinterest. Misjudged posture can make you seem inattentive or aggressive. Sit upright with a subtle forward lean during active listening. • Nodding shows agreement, attentiveness, and engagement. However, Excessive nodding may seem insincere or eager to please. Nod occasionally and naturally. Pair it with eye contact. • Head tilted slightly is a sign of interest, empathy, and engagement. It encourages openness and active listening. Use when listening or acknowledging, combined with eye contact and nodding."
            },
        
        "leg_crossed_count": {
            "title":"Legs Crossed",
            "text": "Sudden Interlocking of Legs or Ankles is a sign of emotional withdrawal or insecurity.  Suggests that the candidate is not fully engaged. Sit with legs naturally open or parallel. Practice with mock interviews. Sitting with Legs Slightly Open Indicates comfort, confidence, and openness.it is a Positive posture when not exaggerated. Shows grounded, relaxed energy. Practice sitting with knees slightly apart, feet flat on the ground. Avoid overly wide stances. Legs or Knees Tightly Held Together- Signals tension, anxiety, or over-control.May appear rigid or overly cautious. Loosen body posture and combine with slow breathing to release tension. Furthermore, Sitting with Legs Crossed, Foot Kicking signals Boredom, agitation, or passive disengagement. Suggests lack of interest or impatience. Keep both feet flat and stable. Shift posture only when necessary.",
            "image": "leg_crossed_count.png"
            },
        
        "leg_bouncing_count": {
            "title":"Leg swaying",
            "text": "Rocking, swaying, or aimless walking shows nervousness or low confidence. Too much movement distracts the audience, pulls attention away from the message, and reduces credibility. Practice grounding your stance with feet shoulder-width apart. Use walking sparingly and intentionally.",
            "image": "leg_bouncing_count.png"
            },
        
        "hand_on_hip_count": {
            "title":"Hands on Hip",
            "text": "Hands or thumbs in pockets Indicates low confidence, passivity, or social discomfort. May diminish your perceived authority or energy. Avoid hiding hands. Keep them relaxed and visible.",
            "image": "hands_on_hip_count.png",
            "additional": "While sitting, the ankle lock gesture is a negative or defensive attitude. It is often combined with clenched fists resting on the knees or with the hands tightly gripping the arms of the chair (male version). The female version varies slightly; the knees are held together, the feet may be to one side, and the hands rest side by side or one on top of the other, resting on the upper legs. The gesture is one of holding back a negative attitude, emotion, nervousness, or fear. If feet are withdrawn under a chair, it shows that the person also has a withdrawn attitude. However, If the legs are outstretched and loosely crossed at the ankle, it usually signals that the person is at ease and comfortable. Placing both feet on the ground with a standard gap between them is the most basic normal position. It serves as a neutral but powerful starting point."
            },
        
        "hands_outside_gesture_box_count": {
            'title':"Purposeful Gestures",
            "text":"Fidgety, aimless, or exaggerated gestures or movements distract rather than reinforce speech. It reduces clarity and credibility. Audience focus shifts from the message to body movements. Keep gestures within the “gesture box” (shoulders to hips). Purposeful, deliberate, and controlled gestures reinforce speech, improve retention, and enhance engagement. Practice emphasizing key words with firm, intentional gestures.",
            "image": "hands_outside_gesture_box.png",
            "additional": "Fondling or twirling hair shows nervousness or lack of self-control. It reduces professionalism and suggests immaturity. Keep your hands away from your face during conversations. The foot lock (while sitting and standing) - This gesture is almost exclusively used by women. The top of one foot locks around the other leg to reinforce a defensive attitude. This position is common to shy or timid women."
            },
        
        "hands_clenched_count": {
            'title':"Hands Clenched",
            'text': "Hands Clenched Together indicate tension, restraint, or anxiety. Often seen as a defensive or nervous gesture.  May reduce perceived confidence and affect message delivery. While sitting, rest your hands openly on your lap or the table. Practice calm breathing to ease hand tension. While standing, rest them lightly at your sides, keep your hands visible, and use purposeful gestures.",
            "image": "hands_clenched.png",
            "additional": "Palm-down gestures convey control, authority, and decisiveness. Can emphasize a point well when used strategically. Use sparingly to convey assertiveness. Balance with open-handed gestures.Palm-up gestures are inviting, collaborative, and inclusive gestures. Encourages participation or agreement. Use when asking questions or offering an opinion to show openness."
            },
        
        "hands_behind_back_count": {
            "title":"Hands Tied Back",
            "text": "In a professional setting, tying hands behind the back is often interpreted as trying to suppress nervous energy and a lack of openness or emotional distancing. It may also signal apprehension, frustration, or discomfort, especially when paired with tense facial expressions or stiff posture. It reduces visibility of the hands, which are key to expressing engagement and sincerity, and may also undermine perceptions of confidence and openness. Avoid holding your hands behind your back during interviews or public speaking. Instead, aim for open, natural gestures in front of your body (within the “gesture box” from shoulders to waist). If you’re unsure what to do with your hands, rest them lightly at your sides or let them gesture naturally when speaking.  ",
            "image": "hands_behind_back.png"
            },
        
        "hands_in_pockets_count": {
            "title":"Hands in Pockets",
            "text": "Hands in pockets suggest low self-confidence, especially in formal or male candidates. Reduces credibility and signals disengagement. Keep hands fully visible. Practice confident, open gestures instead.",
            "image": "hands_in_pockets.jpg",
            "additional": "The buttress stance, i.e, standing with one leg straight and firmly planted, supporting most of the body weight and the other leg bent slightly, often at ease or off to the side, signals readiness or alertness. In some contexts, it can indicate defensiveness or emotional detachment. If combined with closed arms or looking away, it shows lack of engagement. Aim for a balanced stance where both feet are firmly planted shoulder-width apart to project stability and confidence."
            }
        
        }
        
        
        additional_feedback = {
            "open_palms": {
                "title": "Open Palms",
                "text": "Open palms have been associated with honesty, openness, and a willingness to engage. It signals non-aggression and transparency, and is generally well-received by interviewers. It also enhances trust and approachability. Use this gesture purposefully when explaining thoughts or emphasizing sincerity. Avoid excessive gesturing that may look unnatural.",    
                "image": "open_palms.png"
            },
            "expressionless": {
                "title": "Being Expressionless or Neutral throughout",
                "text": "Being expressionless or neutral throughout signals a lack of emotional connection. It makes the delivery flat. The audience may question interest or enthusiasm. Practice facial cues such as smiling, eyebrow raises, and head nods.",
                "image": "expressionless.png",
                "additional": "Chin Tucked In / Nose Down Shows insecurity, anxiety, or withdrawal. Reduces authority and engagement. Keep your chin parallel to the floor.Chin Out / Nose Up - Sign of confidence and assurance. Positive unless exaggerated (may appear arrogant). Use a confident chin posture, especially while entering the room or introducing yourself. "
            },
            "not_smiling": {
                "title": "Not Smiling or a very minimal smile",
                "text": "Not smiling or a very minimal smile throughout the interaction makes one appear stiff, overly serious, or disengaged. The lack of smiling can make the speaker seem unapproachable, nervous, or uninterested. It may lead the interviewer or audience to perceive a lack of enthusiasm or confidence. Practice smiling naturally during greetings, introductions, and when making positive or polite statements. ",
                "image": "not_smiling.png", 
                "additional": "Nose Crinkling suggests dislike, disgust, or disagreement. Can give away negative reactions unintentionally. Keep your face relaxed. Avoid micro-expressions that contradict your spoken message."
                },
            "purposeful_gestures": {
                "title": "Purposeful Gestures",
                "text": "Purposeful gestures are deliberate, meaningful movements that emphasize key points, reinforce clarity and confidence, and match the tone and content of what is being said.Examples include counting on fingers when listing, using open palms when explaining, gesturing upward when describing growth or improvement, etc. When gestures are natural and well-timed, they project confidence, clarity, and engagement. They help the interviewer or the audience better understand and remember what you say. Lack of gestures or overly rigid hands may suggest nervousness or lack of expression. Overly random or excessive gestures can distract and appear unprofessional.Use gestures to support, not replace, your speech. Keep them within the “gesture box” (shoulders to hips). Practice emphasizing key words or transitions with subtle hand movement. Avoid fidgeting, repetitive gestures, or blocking your body. When unsure, open palms and light hand movement are safe, confident defaults.",
                "image": "purposeful_gestures.png",
                "additional": "Mirroring the interviewer’s gestures indicates rapport, agreement, and empathy when done subtly. It creates a positive connection when natural. Mirror body language subtly and only when it feels organic. Avoid forced imitation."
                }, 
            "shrugging":{
                "title": "Shrugging Shoulders",
                "text" :  "Shrugging Shoulders indicates helplessness or uncertainty. It weakens verbal responses and suggests indecision. Avoid shrugging. Instead, clarify with verbal expressions and a confident tone.",
                "image": "shrugging.png",            
            },
            "pursedlips":{
                "title": "Pursed Lips / Lip Biting / Lip Licking",
                "text": "Pursed lips indicate disapproval or disagreement; Biting or licking signals anxiety or uncertainty. Can distract or signal unspoken discomfort. Practice keeping a relaxed, neutral mouth posture.",
                "image": "pursedlips.png",
                "additional": "Tongue Jutting suggests guilt, mischief, or a social faux pas. Highly inappropriate and immature in interviews. Become aware of this via video feedback and train yourself to maintain a composed facial posture."
                
            }
        }


        doc = Document()
        doc.add_heading('Personalized Feedback', level=1)
        
        doc.add_heading("Body Language Score Table", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Body language cues'
        hdr_cells[1].text = 'Observation'
        hdr_cells[2].text = 'Score'
        hdr_cells[3].text = 'Explanation'
        rubricscores = grade["rubric_scores"]
        print(rubricscores)

        print("[INFO]:creating dictionaries")
        # You can customize these based on your detection/logic
        stance = {
            '1' : {
                "cue": "STANCE(Rocking, Swaying, Movement)",
                "observation": "Rocking/swaying more than 5 times",
                "score": rubricscores['stance'],
                "explanation": "Feet not firmly planted, a lot of rocking, swaying, or walking that detracted from message. Throughout the session, observer was focused on stance.."
            },
            '3' : {
                "cue": "STANCE(Rocking, Swaying, Movement)",
                "observation": "Rocking/swaying 3–5 times",
                "score": rubricscores['stance'],
                "explanation": "Occurrences of a neutral stance. Feet mostly planted, some movement, mostly faced audience, and occasional rocking or swaying. Throughout session, observer focused more on message than stance."
            },
            '5' : {
                "cue": "STANCE(Rocking, Swaying, Movement)",
                "observation": "Good stance",
                "score": rubricscores['stance'],
                "explanation": "Adapts a neutral stance. Feet firmly planted, walking did not detract from message, turned to audience, and minimal swaying, throughout session.",
        }
        }
        print(stance)
        
        eye_contact = {
            '1' : {
                "cue": "EYE CONTACT",
                "observation": "Limited eye contact",
                "score": rubricscores['eye_contact'],
                "explanation":"Consistently did not make eye contact with audience members. Constant glancing at ceiling or the floor a majority of the time. Very limited eye contact."
            },
            '3' : {
                "cue": "EYE CONTACT",
                "observation": "maintained eye contact 50-70% times",
                "score": rubricscores['eye_contact'],
                "explanation": "Inconsistently made eye contact with audience members. Frequent glancing for a duration of 1 to 3 s to ceiling or floor or side throughout session (e.g., like a ping pong ball). Intermittent eye contact with audience members. Overall, appropriate eye contact did not occur a majority of the time."
            },
            '5' : {
                "cue": "EYE CONTACT",
                "observation": "Maintains eye contact",
                "score": rubricscores['eye_contact'],
                "explanation": "Consistently and appropriately made eye contact with audience members. Occasional glancing for a duration of 1 to 3 s to ceiling or floor is appropriate throughout entire session. Overall, appropriate eye contact occurred a majority of the time."
        }
        }
        print(eye_contact)
        
        gestures = {
            '1' : {
                "cue": "PURPOSEFUL GESTURES",
                "observation": "No major gestures",
                "score": rubricscores['gestures'],
                "explanation": "Gestures: Extraneous, distracting, or ineffective arm movements that detracted from message. A majority of gestures or gestures/movements were not used to highlight key words and overshadowed purposeful gestures. Observer focused on these extraneous movements (fidgeting, distracting, or used to think rather than highlight)"
            },
            '3' : {
                "cue": "PURPOSEFUL GESTURES",
                "observation": "Zero to one purposeful Gestures",
                "score": rubricscores['gestures'],
                "explanation": "Zero to one purposeful gestures required made between hips and shoulders with firm movement to highlight important word or aspect of message. Instances of “talking” with hands or “weak” gesture attempts (not in hip/shoulder zone)."
            },
            '5' : {
                "cue": "PURPOSEFUL GESTURES",
                "observation": "Performed more than one gesture",
                "score": rubricscores['gestures'],
                "explanation": "At least two purposeful gestures made between hips and shoulders with firm movement to highlight important word or aspect of message. Gestures were purposeful and may have been accompanied by minimal “talking” with hands."
        }
        }
        print(gestures)
        facial_expression = {
            '1' : {
                "cue": "FACIAL EXPRESSIONS",
                "observation": "Neutral expression",
                "score": rubricscores['facial_expression'],
                "explanation": " No facial expressions observed throughout session (raised eyebrows or furrowed brow, wider eyes, smiling, deliberate head movement). Smile at end of exchange does not count."
            },
            '3' : {
                "cue": "FACIAL EXPRESSIONS",
                "observation": "Some expressions were observed",
                "score": rubricscores['facial_expression'],
                "explanation": "Occurrences of facial expressions were intermittent over the session (raised eyebrows or furrowed brow, wider eyes, smiling, deliberate head movement). Possible instances of nonpurposeful distracting facial movements that were repetitive. Smile at end of exchange does not count."
            },
            '5' : {
                "cue": "FACIAL EXPRESSIONS",
                "observation": "Facial expressions were Consistent",
                "score": rubricscores['facial_expression'],
                "explanation": "Four or more instances of facial expressions across two or more responses (raised eyebrows or furrowed brow, wider eyes, smiling, deliberate head movement). No distracting movements. Smile at end of exchange does not count."
            }
        } 
        print(facial_expression)
        
        print("[INFO]:Dictionaries created")
        
        stance_score=str(rubricscores['stance'])
        print(stance_score)
        eyecontact_score=str(rubricscores['eye_contact'])
        gesture_score=str(rubricscores['gestures'])
        facial_score=str(rubricscores['facial_expression'])
        print(facial_score)
        
        print("[INFO]:Made dict acc to score")

        stancee=stance[stance_score]
        eyee=eye_contact[eyecontact_score]
        gesturee=gestures[gesture_score]
        faciall=facial_expression[facial_score] 

        print("[INFO]:Collected details for table")

        
        row_cells = table.add_row().cells
        row_cells[0].text = stancee["cue"].replace('\xa0', ' ')
        row_cells[1].text = stancee["observation"].replace('\xa0', ' ')
        row_cells[2].text = str(stancee["score"])
        row_cells[3].text = stancee["explanation"].replace('\xa0', ' ')

        row_cells = table.add_row().cells
        row_cells[0].text = eyee["cue"].replace('\xa0', ' ')
        row_cells[1].text = eyee["observation"].replace('\xa0', ' ')
        row_cells[2].text = str(eyee["score"])
        row_cells[3].text = eyee["explanation"].replace('\xa0', ' ')

        row_cells = table.add_row().cells
        row_cells[0].text = gesturee["cue"].replace('\xa0', ' ')
        row_cells[1].text = gesturee["observation"].replace('\xa0', ' ')
        row_cells[2].text = str(gesturee["score"])
        row_cells[3].text = gesturee["explanation"].replace('\xa0', ' ')

        row_cells = table.add_row().cells
        row_cells[0].text = faciall["cue"].replace('\xa0', ' ')
        row_cells[1].text = faciall["observation"].replace('\xa0', ' ')
        row_cells[2].text = str(faciall["score"])
        row_cells[3].text = faciall["explanation"].replace('\xa0', ' ')

        

        print("[INFO]:Table added")
        
        
        doc.add_paragraph(f"Report:)")
        doc.add_paragraph(f"Duration: {report['duration_sec']} seconds")
        doc.add_paragraph(f"Final BPM: {report.get('final_bpm', 'N/A')}")
        doc.add_paragraph(f"Strictness Level: {strictness}")
        doc.add_paragraph("Feedback Summary:")
        doc.add_paragraph(f"Eye Contact Breaks: {normalized_report['eye_contact_breaks']:.2f} per minute")
        doc.add_paragraph(f"Total Blinks: {normalized_report['total_blinks']:.2f} per minute")
        doc.add_paragraph(f"Mouth Touches: {normalized_report['mouth_touch_count']:.2f} per minute")
        doc.add_paragraph(f"Nose Touches: {normalized_report['nose_touch_count']:.2f} per minute")
        doc.add_paragraph(f"Eye Touches: {normalized_report['eye_touch_count']:.2f} per minute")
        doc.add_paragraph(f"Ear Touches: {normalized_report['ear_touch_count']:.2f} per minute")
        doc.add_paragraph(f"Neck Touches: {normalized_report['neck_touch_count']:.2f} per minute")
        doc.add_paragraph(f"Arms Crossed for 3 Seconds: {normalized_report['arms_crossed_for_3_sec_count']:.2f} per minute")
        doc.add_paragraph(f"Slouching: {normalized_report['slouching']:.2f} per minute")
        doc.add_paragraph(f"Hands Clenched: {normalized_report['hands_clenched_count']:.2f} per minute")
        doc.add_paragraph(f"Hands Behind Back: {normalized_report['hands_behind_back_count']:.2f} per minute")
        doc.add_paragraph(f"Hands in Pockets: {normalized_report['hands_in_pockets_count']:.2f} per minute")
        doc.add_paragraph(f"Hands Outside Gesture Box: {normalized_report['hands_outside_gesture_box_count']:.2f} per minute")
        doc.add_paragraph(f"Legs Crossed: {normalized_report['leg_crossed_count']:.2f} per minute")
        doc.add_paragraph(f"Leg Bouncing: {normalized_report['leg_bouncing_count']:.2f} per minute")
        doc.add_paragraph(f"Hand on Hip: {normalized_report['hand_on_hip_count']:.2f} per minute")


        # Detailed feedback
        print("[INFO] Adding detailed feedback...")
        for gesture, value in normalized_report.items():
            if is_abnormal(gesture, value,abnormal_thresholds):
                if gesture in gesture_feedback:
                    feedback = gesture_feedback[gesture]
                    doc.add_heading(feedback["title"], level=2)
                    doc.add_paragraph(f"Observed: {round(value, 2)} per minute")
                    doc.add_paragraph(feedback["text"])

                    if "image" in feedback:
                        image_path = os.path.join("images", feedback["image"])
                        print(f"[INFO] Attempting to load image for {gesture}: {image_path}")
                        if os.path.exists(image_path):
                            doc.add_picture(image_path, width=Inches(2))
                            print(f"[SUCCESS] Image added: {image_path}")
                        else:
                            print(f"[WARNING] Image not found: {image_path}")
                    if "additional" in feedback:
                        doc.add_heading("Additional Notes", level=3)
                        doc.add_paragraph(feedback["additional"])
                else:
                    print(f"[WARNING] No feedback found for abnormal gesture: {gesture}")
            else:
                print(f"[INFO] {gesture} is within normal range.")
                print("Adding hands behind back")        
        if normalized_report["hands_in_pockets_count"] >= abnormal_thresholds.get("hands_in_pockets_count", 3):
                    doc.add_heading("Hands behind Back", level=2)
                    feed=gesture_feedback["hands_behind_back_count"]
                    doc.add_paragraph(feed["text"])
                    if "image" in feed:
                        image_path = os.path.join("images", feed["image"])
                        print(f"[INFO] Attempting to load image for hands behind back: {image_path}")
                        if os.path.exists(image_path):
                            doc.add_picture(image_path, width=Inches(2))
                            print(f"[SUCCESS] Image added: {image_path}")
                        else:
                            print(f"[WARNING] Image not found: {image_path}")
                    
                    if "additional" in feed:
                        doc.add_heading("Additional Notes", level=3)
                        doc.add_paragraph(feed["additional"])
        print("Added hands behind back")            
        # Additional feedback
        print("[INFO] Adding additional feedback...")
        for key, feedback in additional_feedback.items():
            doc.add_heading(feedback["title"], level=2)
            doc.add_paragraph(feedback["text"])
            if "image" in feedback:
                image_path = os.path.join("images", feedback["image"])
                print(f"[INFO] Attempting to load image for {key}: {image_path}")
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(2))
                    print(f"[SUCCESS] Image added: {image_path}")
                else:
                    print(f"[WARNING] Image not found: {image_path}")
            if "additional" in feedback:
                doc.add_heading("Additional Notes", level=3)
                doc.add_paragraph(feedback["additional"])

        output_path = get_next_feedback_filename()
        doc.save(output_path)
        print(f"[SUCCESS] Feedback document saved to {output_path}")

    except Exception as e:
        print(f"[ERROR] Failed to generate feedback document: {str(e)}")
        
        
def generate_feedback_folder(report, emotion_stats , grade ,file_name ,abnormal_thresholds,strictness=1):
    print("[INFO] Starting feedback document generation...")

    try:
        # Normalize based on strictness level
        if strictness == 2:
            duration_unit = max(1, report["duration_sec"] / 30)
        elif strictness == 3:
            duration_unit = max(1, report["duration_sec"] / 15)
        elif strictness == 4:
            duration_unit = max(1, report["duration_sec"] / 60)
        else:
            duration_unit = 1
        print(f"[INFO] Duration unit calculated: {duration_unit:.2f} seconds")

        norm = lambda count: count / duration_unit

        # Normalize behavior counts
        print("[INFO] Normalizing behavior metrics...")
        normalized_report = {
            "eye_contact_breaks": norm(report["eye_contact_breaks"]),
            "total_blinks": norm(report["total_blinks"]),
            "mouth_touch_count": norm(report["mouth_touch_count"]),
            "nose_touch_count": norm(report["nose_touch_count"]),
            "eye_touch_count": norm(report["eye_touch_count"]),
            "ear_touch_count": norm(report["ear_touch_count"]),
            "neck_touch_count": norm(report["neck_touch_count"]),
            "arms_crossed_for_3_sec_count": norm(report["arms_crossed_for_3_sec_count"]),
            "slouching": norm(report["slouching"]),
            "leg_crossed_count": norm(report["leg_crossed_count"]),
            "leg_bouncing_count": norm(report["leg_bouncing_count"]),
            "hand_on_hip_count": norm(report.get("hand_on_hip_count", 0)),
            "final_bpm": report.get("final_bpm", 0)
        }

        print("[INFO] Normalization complete.")

        # Load gesture feedback
        print("[INFO] Preparing feedback descriptions...")
        gesture_feedback = {
        "eye_contact_breaks": {
            "title":"Eye Contact",
            "text": "Avoids eye contact, Glancing at ceiling or floor reflects discomfort or disengagement. Audience feels disconnected; speaker may seem uncertain or underprepared. Practice maintaining soft, steady eye contact with various parts of the audience. ",
            "image": "eye_contact_breaks.png"
            },
        "total_blinks": {
            "title":"Rapid Blinking",
            "text": "Rapid or repeated blinking (more than 20–30 blinks per minute) is often a sign of nervousness, stress, or cognitive overload. It may also indicate discomfort with eye contact or dishonesty, depending on the context. The audience or interviewer may interpret the behavior as a lack of confidence, anxiety, or even evasiveness. It disrupts engagement and reduces perceived trustworthiness. Practice calming techniques such as deep breathing and speaking slower. Improve eye contact comfort by rehearsing with a friend, mirror, or camera. Reduce stress to regulate blinking naturally.",
            "image": "total_blinks.png"
            },
        
        "mouth_touch_count": {
            "title":"Mouth Touching",
            "text" : "Touching the face, mouth, nose, eyes, ears, and neck are soothing gestures that indicate internal discomfort, anxiety, or deception. These displays are common under stress. It can be distracting and potentially lower credibility. Reduces confidence in what’s being said. Increase awareness of this habit. Keep your hands visible and use purposeful gestures instead. Neck / Face Rubbing or Scratching is a Sign of doubt, insecurity, or emotional discomfort. It may weaken trust and show hesitation. Replace with neutral hand positions and affirming verbal cues. ",
            "image": "mouth_touch_count.png",
            "additional": "Fingers in Mouth / Nail Biting are strong pacifying behaviours, often a response to anxiety or lack of self-control. Appears unprofessional and immature. Therefore, it should be avoided. Men tend to massage or stroke their necks to pacify distress. This area is rich with nerves, including the vagus nerve, which, when massaged, will slow down the heart rate. Even a brief touch of the neck will serve to assuage anxiety or discomfort. Neck touching or massaging is a powerful and universal stress reliever and pacifier. It should be avoided."
            },
        
        "nose_touch_count": {
            "title":"Nose Touching",
            "text" : "Touching Face Mouth, Nose, Eyes, Ears and Neck are Soothing gesture that indicates internal discomfort. anxiety, or deception. Common under stress. Can be distracting and signals discomfort, potentially lowering credibility. Reduces confidence in what’s being said. Increase awareness of this habit. Keep hands visible and use purposeful gestures instead.",
            "image": "nose_touch_count.png"
            },
        
        "eye_touch_count": {
            "title":"Eye Touching",
            "text" : "Touching Eyes is a common pacifying behavior, often a response to anxiety or discomfort. It can be distracting and may signal nervousness or dishonesty. Avoid touching your eyes during presentations or interviews to maintain professionalism and credibility. Also, Eyebrow Movement Adds emotional depth; raised brows shows interest and furrowed coneys concern or confusion. Can emphasize or distort your message if mismatched. Use eyebrows expressively but in sync with tone and content.",
            "image": "eye_touch_count.png"
            },
        
        "ear_touch_count": {
            "title":"Ear Touching",
            "text" : "Touching your ears is a common self-soothing gesture, typically triggered by anxiety or discomfort. However, it can be distracting and may unintentionally signal nervousness or a lack of honesty. To maintain a professional and confident impression during presentations or interviews, it's best to avoid this behavior.",
            "image": "ear_touch_count.png"
            },
        
        "neck_touch_count": {
            "title":"Neck Touching",
            "text" : "Scratching the neck is typically interpreted as a pacifying behavior, often linked to self-doubt, stress, or internal conflict. It's a subconscious gesture that may indicate the person is feeling uncertain, uncomfortable, or lacking confidence in what they're saying. In high-stakes settings like interviews or public speaking, this gesture can subtly undermine your message, suggesting hesitation or insincerity. For a more confident and composed presence, it's advisable to become aware of and minimize such self-touch behaviors.",
            "image": "neck_touch_count.png"
            },
        
        "arms_crossed_for_3_sec_count": {
            "title":"Arms Crossed",
            "text": "Both arms are folded together across the chest is a Defensive or protective gesture. Reduces approachability and may signal emotional distancing, a negative attitude, shyness or nervousness. Use open posture with relaxed arms. Crossed Arms with Hands Gripping Arms also conveys Strong emotional discomfort or internal resistance." ,
            "image": "arms_crossed.png"
            },
        
        "slouching": {
            "title":"Slouching",
            "text": "Slouching Indicates lack of composure, energy, or seriousness. Projects low confidence and poor presence. Practice upright posture with an open chest, feet grounded, and shoulders relaxed.",
            "image": "slouching.png",
            "additional": " • A slight forward lean signals engagement. However, leaning back indicates disinterest. Misjudged posture can make you seem inattentive or aggressive. Sit upright with a subtle forward lean during active listening. • Nodding shows agreement, attentiveness, and engagement. However, Excessive nodding may seem insincere or eager to please. Nod occasionally and naturally. Pair it with eye contact. • Head tilted slightly is a sign of interest, empathy, and engagement. It encourages openness and active listening. Use when listening or acknowledging, combined with eye contact and nodding."
            },
        
        "leg_crossed_count": {
            "title":"Legs Crossed",
            "text": "Sudden Interlocking of Legs or Ankles is a sign of emotional withdrawal or insecurity.  Suggests that the candidate is not fully engaged. Sit with legs naturally open or parallel. Practice with mock interviews. Sitting with Legs Slightly Open Indicates comfort, confidence, and openness.it is a Positive posture when not exaggerated. Shows grounded, relaxed energy. Practice sitting with knees slightly apart, feet flat on the ground. Avoid overly wide stances. Legs or Knees Tightly Held Together- Signals tension, anxiety, or over-control.May appear rigid or overly cautious. Loosen body posture and combine with slow breathing to release tension. Furthermore, Sitting with Legs Crossed, Foot Kicking signals Boredom, agitation, or passive disengagement. Suggests lack of interest or impatience. Keep both feet flat and stable. Shift posture only when necessary.",
            "image": "leg_crossed_count.png"
            },
        
        "leg_bouncing_count": {
            "title":"Leg swaying",
            "text": "Rocking, swaying, or aimless walking shows nervousness or low confidence. Too much movement distracts the audience, pulls attention away from the message, and reduces credibility. Practice grounding your stance with feet shoulder-width apart. Use walking sparingly and intentionally.",
            "image": "leg_bouncing_count.png"
            },
        
        "hand_on_hip_count": {
            "title":"Hands on Hip",
            "text": "Hands or thumbs in pockets Indicates low confidence, passivity, or social discomfort. May diminish your perceived authority or energy. Avoid hiding hands. Keep them relaxed and visible.",
            "image": "hands_on_hip_count.png",
            "additional": "While sitting, the ankle lock gesture is a negative or defensive attitude. It is often combined with clenched fists resting on the knees or with the hands tightly gripping the arms of the chair (male version). The female version varies slightly; the knees are held together, the feet may be to one side, and the hands rest side by side or one on top of the other, resting on the upper legs. The gesture is one of holding back a negative attitude, emotion, nervousness, or fear. If feet are withdrawn under a chair, it shows that the person also has a withdrawn attitude. However, If the legs are outstretched and loosely crossed at the ankle, it usually signals that the person is at ease and comfortable. Placing both feet on the ground with a standard gap between them is the most basic normal position. It serves as a neutral but powerful starting point."
            },
        
        "hands_outside_gesture_box_count": {
            'title':"Purposeful Gestures",
            "text":"Fidgety, aimless, or exaggerated gestures or movements distract rather than reinforce speech. It reduces clarity and credibility. Audience focus shifts from the message to body movements. Keep gestures within the “gesture box” (shoulders to hips). Purposeful, deliberate, and controlled gestures reinforce speech, improve retention, and enhance engagement. Practice emphasizing key words with firm, intentional gestures.",
            "image": "hands_outside_gesture_box.png",
            "additional": "Fondling or twirling hair shows nervousness or lack of self-control. It reduces professionalism and suggests immaturity. Keep your hands away from your face during conversations. The foot lock (while sitting and standing) - This gesture is almost exclusively used by women. The top of one foot locks around the other leg to reinforce a defensive attitude. This position is common to shy or timid women."
            },
        
        "hands_clenched_count": {
            'title':"Hands Clenched",
            'text': "Hands Clenched Together indicate tension, restraint, or anxiety. Often seen as a defensive or nervous gesture.  May reduce perceived confidence and affect message delivery. While sitting, rest your hands openly on your lap or the table. Practice calm breathing to ease hand tension. While standing, rest them lightly at your sides, keep your hands visible, and use purposeful gestures.",
            "image": "hands_clenched.png",
            "additional": "Palm-down gestures convey control, authority, and decisiveness. Can emphasize a point well when used strategically. Use sparingly to convey assertiveness. Balance with open-handed gestures.Palm-up gestures are inviting, collaborative, and inclusive gestures. Encourages participation or agreement. Use when asking questions or offering an opinion to show openness."
            },
        
        "hands_behind_back_count": {
            "title":"Hands Tied Back",
            "text": "In a professional setting, tying hands behind the back is often interpreted as trying to suppress nervous energy and a lack of openness or emotional distancing. It may also signal apprehension, frustration, or discomfort, especially when paired with tense facial expressions or stiff posture. It reduces visibility of the hands, which are key to expressing engagement and sincerity, and may also undermine perceptions of confidence and openness. Avoid holding your hands behind your back during interviews or public speaking. Instead, aim for open, natural gestures in front of your body (within the “gesture box” from shoulders to waist). If you’re unsure what to do with your hands, rest them lightly at your sides or let them gesture naturally when speaking.  ",
            "image": "hands_behind_back.png"
            },
        
        "hands_in_pockets_count": {
            "title":"Hands in Pockets",
            "text": "Hands in pockets suggest low self-confidence, especially in formal or male candidates. Reduces credibility and signals disengagement. Keep hands fully visible. Practice confident, open gestures instead.",
            "image": "hands_in_pockets.png",
            "additional": "The buttress stance, i.e, standing with one leg straight and firmly planted, supporting most of the body weight and the other leg bent slightly, often at ease or off to the side, signals readiness or alertness. In some contexts, it can indicate defensiveness or emotional detachment. If combined with closed arms or looking away, it shows lack of engagement. Aim for a balanced stance where both feet are firmly planted shoulder-width apart to project stability and confidence."
            }
        
        }
        
        additional_feedback = {
            "open_palms": {
                "title": "Open Palms",
                "text": "Open palms have been associated with honesty, openness, and a willingness to engage. It signals non-aggression and transparency, and is generally well-received by interviewers. It also enhances trust and approachability. Use this gesture purposefully when explaining thoughts or emphasizing sincerity. Avoid excessive gesturing that may look unnatural.",    
                "image": "open_palms.png"
            },
            "expressionless": {
                "title": "Being Expressionless or Neutral throughout",
                "text": "Being expressionless or neutral throughout signals a lack of emotional connection. It makes the delivery flat. The audience may question interest or enthusiasm. Practice facial cues such as smiling, eyebrow raises, and head nods.",
                "image": "expressionless.png",
                "additional": "Chin Tucked In / Nose Down Shows insecurity, anxiety, or withdrawal. Reduces authority and engagement. Keep your chin parallel to the floor.Chin Out / Nose Up - Sign of confidence and assurance. Positive unless exaggerated (may appear arrogant). Use a confident chin posture, especially while entering the room or introducing yourself. "
            },
            "not_smiling": {
                "title": "Not Smiling or a very minimal smile",
                "text": "Not smiling or a very minimal smile throughout the interaction makes one appear stiff, overly serious, or disengaged. The lack of smiling can make the speaker seem unapproachable, nervous, or uninterested. It may lead the interviewer or audience to perceive a lack of enthusiasm or confidence. Practice smiling naturally during greetings, introductions, and when making positive or polite statements. ",
                "image": "not_smiling.png", 
                "additional": "Nose Crinkling suggests dislike, disgust, or disagreement. Can give away negative reactions unintentionally. Keep your face relaxed. Avoid micro-expressions that contradict your spoken message."
                },
            "purposeful_gestures": {
                "title": "Purposeful Gestures",
                "text": "Purposeful gestures are deliberate, meaningful movements that emphasize key points, reinforce clarity and confidence, and match the tone and content of what is being said.Examples include counting on fingers when listing, using open palms when explaining, gesturing upward when describing growth or improvement, etc. When gestures are natural and well-timed, they project confidence, clarity, and engagement. They help the interviewer or the audience better understand and remember what you say. Lack of gestures or overly rigid hands may suggest nervousness or lack of expression. Overly random or excessive gestures can distract and appear unprofessional.Use gestures to support, not replace, your speech. Keep them within the “gesture box” (shoulders to hips). Practice emphasizing key words or transitions with subtle hand movement. Avoid fidgeting, repetitive gestures, or blocking your body. When unsure, open palms and light hand movement are safe, confident defaults.",
                "image": "purposeful_gestures.png",
                "additional": "Mirroring the interviewer’s gestures indicates rapport, agreement, and empathy when done subtly. It creates a positive connection when natural. Mirror body language subtly and only when it feels organic. Avoid forced imitation."
                }, 
            "shrugging":{
                "title": "Shrugging Shoulders",
                "text" :  "Shrugging Shoulders indicates helplessness or uncertainty. It weakens verbal responses and suggests indecision. Avoid shrugging. Instead, clarify with verbal expressions and a confident tone.",
                "image": "shrugging.png",            
            },
            "pursedlips":{
                "title": "Pursed Lips / Lip Biting / Lip Licking",
                "text": "Pursed lips indicate disapproval or disagreement; Biting or licking signals anxiety or uncertainty. Can distract or signal unspoken discomfort. Practice keeping a relaxed, neutral mouth posture.",
                "image": "pursedlips.png",
                "additional": "Tongue Jutting suggests guilt, mischief, or a social faux pas. Highly inappropriate and immature in interviews. Become aware of this via video feedback and train yourself to maintain a composed facial posture."
                
            }
        }


        doc = Document()
        doc.add_heading('Personalized Feedback', level=1)
        
                # Add Body Language Cue Table
        doc.add_heading("Body Language Score Table", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Body language cues'
        hdr_cells[1].text = 'Observation'
        hdr_cells[2].text = 'Score'
        hdr_cells[3].text = 'Explanation'
        rubricscores = grade["rubric_scores"]
        print(rubricscores)

        print("[INFO]:creating dictionaries")
        # You can customize these based on your detection/logic
        stance = {
            '1' : {
                "cue": "STANCE(Rocking, Swaying, Movement)",
                "observation": "Rocking/swaying more than 5 times",
                "score": rubricscores['stance'],
                "explanation": "Feet not firmly planted, a lot of rocking, swaying, or walking that detracted from message. Throughout the session, observer was focused on stance.."
            },
            '3' : {
                "cue": "STANCE(Rocking, Swaying, Movement)",
                "observation": "Rocking/swaying 3–5 times",
                "score": rubricscores['stance'],
                "explanation": "Occurrences of a neutral stance. Feet mostly planted, some movement, mostly faced audience, and occasional rocking or swaying. Throughout session, observer focused more on message than stance."
            },
            '5' : {
                "cue": "STANCE(Rocking, Swaying, Movement)",
                "observation": "Good stance",
                "score": rubricscores['stance'],
                "explanation": "Adapts a neutral stance. Feet firmly planted, walking did not detract from message, turned to audience, and minimal swaying, throughout session.",
        }
        }
        print(stance)
        
        eye_contact = {
            '1' : {
                "cue": "EYE CONTACT",
                "observation": "Limited eye contact",
                "score": rubricscores['eye_contact'],
                "explanation":"Consistently did not make eye contact with audience members. Constant glancing at ceiling or the floor a majority of the time. Very limited eye contact."
            },
            '3' : {
                "cue": "EYE CONTACT",
                "observation": "maintained eye contact 50-70% times",
                "score": rubricscores['eye_contact'],
                "explanation": "Inconsistently made eye contact with audience members. Frequent glancing for a duration of 1 to 3 s to ceiling or floor or side throughout session (e.g., like a ping pong ball). Intermittent eye contact with audience members. Overall, appropriate eye contact did not occur a majority of the time."
            },
            '5' : {
                "cue": "EYE CONTACT",
                "observation": "Maintains eye contact",
                "score": rubricscores['eye_contact'],
                "explanation": "Consistently and appropriately made eye contact with audience members. Occasional glancing for a duration of 1 to 3 s to ceiling or floor is appropriate throughout entire session. Overall, appropriate eye contact occurred a majority of the time."
        }
        }
        print(eye_contact)
        
        gestures = {
            '1' : {
                "cue": "PURPOSEFUL GESTURES",
                "observation": "No major gestures",
                "score": rubricscores['gestures'],
                "explanation": "Gestures: Extraneous, distracting, or ineffective arm movements that detracted from message. A majority of gestures or gestures/movements were not used to highlight key words and overshadowed purposeful gestures. Observer focused on these extraneous movements (fidgeting, distracting, or used to think rather than highlight)"
            },
            '3' : {
                "cue": "PURPOSEFUL GESTURES",
                "observation": "Zero to one purposeful Gestures",
                "score": rubricscores['gestures'],
                "explanation": "Zero to one purposeful gestures required made between hips and shoulders with firm movement to highlight important word or aspect of message. Instances of “talking” with hands or “weak” gesture attempts (not in hip/shoulder zone)."
            },
            '5' : {
                "cue": "PURPOSEFUL GESTURES",
                "observation": "Performed more than one gesture",
                "score": rubricscores['gestures'],
                "explanation": "At least two purposeful gestures made between hips and shoulders with firm movement to highlight important word or aspect of message. Gestures were purposeful and may have been accompanied by minimal “talking” with hands."
        }
        }
        print(gestures)
        facial_expression = {
            '1' : {
                "cue": "FACIAL EXPRESSIONS",
                "observation": "Neutral expression",
                "score": rubricscores['facial_expression'],
                "explanation": " No facial expressions observed throughout session (raised eyebrows or furrowed brow, wider eyes, smiling, deliberate head movement). Smile at end of exchange does not count."
            },
            '3' : {
                "cue": "FACIAL EXPRESSIONS",
                "observation": "Some expressions were observed",
                "score": rubricscores['facial_expression'],
                "explanation": "Occurrences of facial expressions were intermittent over the session (raised eyebrows or furrowed brow, wider eyes, smiling, deliberate head movement). Possible instances of nonpurposeful distracting facial movements that were repetitive. Smile at end of exchange does not count."
            },
            '5' : {
                "cue": "FACIAL EXPRESSIONS",
                "observation": "Facial expressions were Consistent",
                "score": rubricscores['facial_expression'],
                "explanation": "Four or more instances of facial expressions across two or more responses (raised eyebrows or furrowed brow, wider eyes, smiling, deliberate head movement). No distracting movements. Smile at end of exchange does not count."
            }
        } 
        print(facial_expression)
        
        print("[INFO]:Dictionaries created")
        
        stance_score=str(rubricscores['stance'])
        print(stance_score)
        eyecontact_score=str(rubricscores['eye_contact'])
        gesture_score=str(rubricscores['gestures'])
        facial_score=str(rubricscores['facial_expression'])
        print(facial_score)
        
        print("[INFO]:Made dict acc to score")

        stancee=stance[stance_score]
        eyee=eye_contact[eyecontact_score]
        gesturee=gestures[gesture_score]
        faciall=facial_expression[facial_score] 

        print("[INFO]:Collected details for table")

        row_cells = table.add_row().cells
        row_cells[0].text = stancee["cue"].replace('\xa0', ' ')
        row_cells[1].text = stancee["observation"].replace('\xa0', ' ')
        row_cells[2].text = str(stancee["score"])
        row_cells[3].text = stancee["explanation"].replace('\xa0', ' ')

        row_cells = table.add_row().cells
        row_cells[0].text = eyee["cue"].replace('\xa0', ' ')
        row_cells[1].text = eyee["observation"].replace('\xa0', ' ')
        row_cells[2].text = str(eyee["score"])
        row_cells[3].text = eyee["explanation"].replace('\xa0', ' ')

        row_cells = table.add_row().cells
        row_cells[0].text = gesturee["cue"].replace('\xa0', ' ')
        row_cells[1].text = gesturee["observation"].replace('\xa0', ' ')
        row_cells[2].text = str(gesturee["score"])
        row_cells[3].text = gesturee["explanation"].replace('\xa0', ' ')

        row_cells = table.add_row().cells
        row_cells[0].text = faciall["cue"].replace('\xa0', ' ')
        row_cells[1].text = faciall["observation"].replace('\xa0', ' ')
        row_cells[2].text = str(faciall["score"])
        row_cells[3].text = faciall["explanation"].replace('\xa0', ' ')
        

        print("[INFO]:Table added")
        
        
        
        
        doc.add_paragraph(f"Report:")
        doc.add_paragraph(f"Duration: {report['duration_sec']} seconds")
        doc.add_paragraph(f"Final BPM: {report.get('final_bpm', 'N/A')}")
        doc.add_paragraph(f"Strictness Level: {strictness}")
        doc.add_paragraph("Feedback Summary:")
        doc.add_paragraph(f"Eye Contact Breaks: {normalized_report['eye_contact_breaks']:.2f} per minute")
        doc.add_paragraph(f"Total Blinks: {normalized_report['total_blinks']:.2f} per minute")
        doc.add_paragraph(f"Mouth Touches: {normalized_report['mouth_touch_count']:.2f} per minute")
        doc.add_paragraph(f"Nose Touches: {normalized_report['nose_touch_count']:.2f} per minute")
        doc.add_paragraph(f"Eye Touches: {normalized_report['eye_touch_count']:.2f} per minute")
        doc.add_paragraph(f"Ear Touches: {normalized_report['ear_touch_count']:.2f} per minute")
        doc.add_paragraph(f"Neck Touches: {normalized_report['neck_touch_count']:.2f} per minute")
        doc.add_paragraph(f"Arms Crossed for 3 Seconds: {normalized_report['arms_crossed_for_3_sec_count']:.2f} per minute")
        doc.add_paragraph(f"Slouching: {normalized_report['slouching']:.2f} per minute")
        doc.add_paragraph(f"Legs Crossed: {normalized_report['leg_crossed_count']:.2f} per minute")
        doc.add_paragraph(f"Leg Bouncing: {normalized_report['leg_bouncing_count']:.2f} per minute")
        doc.add_paragraph(f"Hand on Hip: {normalized_report['hand_on_hip_count']:.2f} per minute")
        doc.add_paragraph(f"Hands Outside Gesture Box: {normalized_report.get('hands_outside_gesture_box_count', 0):.2f} per minute")
        doc.add_paragraph(f"Hands Clenched: {normalized_report.get('hands_clenched_count', 0):.2f} per minute")
        doc.add_paragraph(f"Hands Behind Back: {normalized_report.get('hands_behind_back_count', 0):.2f} per minute")
        # doc.add_paragraph(f"Hands in Pockets: {normalized_report.get('hands_in_pockets_count', 0):.2f} per minute")
        doc.add_paragraph(" ")
        doc.add_paragraph(f"Emotion Stats: {emotion_stats}")
        doc.add_paragraph(f"[SCORE] {grade['rubric_scores']}")
        doc.add_paragraph(f"[RESULT] Total Score: {grade['total_score']}, Average Score: {grade['average_score']} , Rank: {grade['rating']}")


        print("[INFO] Adding detailed feedback...")
        for gesture, value in normalized_report.items():
            if is_abnormal(gesture, value,abnormal_thresholds):
                if gesture in gesture_feedback:
                    feedback = gesture_feedback[gesture]
                    doc.add_heading(feedback["title"], level=2)
                    doc.add_paragraph(f"Observed: {round(value, 2)} per minute")
                    doc.add_paragraph(feedback["text"])

                    if "image" in feedback:
                        image_path = os.path.join("images", feedback["image"])
                        print(f"[INFO] Attempting to load image for {gesture}: {image_path}")
                        if os.path.exists(image_path):
                            doc.add_picture(image_path, width=Inches(2))
                            print(f"[SUCCESS] Image added: {image_path}")
                        else:
                            print(f"[WARNING] Image not found: {image_path}")
                    
                    if "additional" in feedback:
                        doc.add_heading("Additional Tips", level=3)
                        doc.add_paragraph(feedback["additional"])
                else:
                    print(f"[WARNING] No feedback found for abnormal gesture: {gesture}")
            else:
                print(f"[INFO] {gesture} is within normal range.")
        print("Adding hands behind back")        
        # if normalized_report["hands_in_pockets_count"] >= abnormal_thresholds.get("hands_in_pockets_count", 3):
        #             doc.add_heading("Hands behind Back", level=2)
        #             feed=gesture_feedback["hands_behind_back_count"]
        #             doc.add_paragraph(feed["text"])
        #             if "image" in feed:
        #                 image_path = os.path.join("images", feed["image"])
        #                 print(f"[INFO] Attempting to load image for hands behind back: {image_path}")
        #                 if os.path.exists(image_path):
        #                     doc.add_picture(image_path, width=Inches(2))
        #                     print(f"[SUCCESS] Image added: {image_path}")
        #                 else:
        #                     print(f"[WARNING] Image not found: {image_path}")
                    
        #             if "additional" in feed:
        #                 doc.add_heading("Additional Notes", level=3)
        #                 doc.add_paragraph(feed["additional"])
        # print("Added hands behind back")                        
        # Additional feedback
        print("[INFO] Adding additional feedback...")
        for key, feedback in additional_feedback.items():
            doc.add_heading(feedback["title"], level=2)
            doc.add_paragraph(feedback["text"])
            if "image" in feedback:
                image_path = os.path.join("images", feedback["image"])
                print(f"[INFO] Attempting to load image for {key}: {image_path}")
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(2))
                    print(f"[SUCCESS] Image added: {image_path}")
                else:
                    print(f"[WARNING] Image not found: {image_path}")
            if "additional" in feedback:
                doc.add_heading("Additional Notes", level=3)
                doc.add_paragraph(feedback["additional"])
        output_path = file_name
        if not output_path.endswith('.docx'):
            output_path += '.docx'
        doc.save(output_path)
        print(f"[SUCCESS] Feedback document saved to {output_path}")

    except Exception as e:
        print(f"[ERROR] Failed to generate feedback document: {str(e)}")

def is_abnormal(behavior,count,abnormal_threshold):
    # abnormal_threshold = {
    #     "eye_contact_breaks": 4,
    #     "total_blinks": 20,
    #     "mouth_touch_count": 1,
    #     "nose_touch_count": 1,
    #     "eye_touch_count": 1,
    #     "ear_touch_count": 1,
    #     "neck_touch_count": 1,
    #     "arms_crossed_for_3_sec_count": 2,
    #     "slouching": 1,
    #     "leg_crossed_count": 2,
    #     "leg_bouncing_count": 2,
    #     "hand_on_hip_count": 2,
    #     "final_bpm": 20
    # }

    threshold = abnormal_threshold.get(behavior, None)
    if threshold is None:
        print(f"[WARNING] No threshold defined for: {behavior}")
        return False
    return count > threshold

def get_next_feedback_filename(folder="feedback"):
    os.makedirs(folder, exist_ok=True)
    i = 1
    while True:
        filename = f"feedback_{i}.docx"
        filepath = os.path.join(folder, filename)
        if not os.path.exists(filepath):
            print(f"[INFO] Generated filename: {filepath}")
            return filepath
        i += 1
